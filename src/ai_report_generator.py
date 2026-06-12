"""AI 报告生成编排模块。

编排流程:
  1. 打包分析结果为 JSON payload（analysis_packager）
  2. 构建 LLM 系统/用户提示词（llm_prompts）
  3. 调用 LLM 生成报告文本（llm_client）
  4. 将 Markdown 报告转换为 HTML / DOCX

所有统计计算由 pandas/scipy/statsmodels 完成，AI 只负责解释和撰写。

支持：
  - 5 种报告结构（通用调研/学术论文/政务决策/商业分析/课程作业）
  - 4 种写作风格（课程作业风/政务汇报风/学术报告风/商业分析风）
  - 3 种报告长度（简短版/标准版/详细版）
  - 5 种 HTML 导出主题
"""

import io
import logging
import re
from pathlib import Path
from typing import Optional, Dict, Any, List

import pandas as pd

from src.analysis_packager import build_analysis_payload, to_json_payload
from src.llm_prompts import build_ai_report_prompt
from src.llm_client import call_llm
from src.html_report_templates import render_html_report, get_available_themes
from src.report_context import build_variable_summary, build_key_findings_summary

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "outputs"


# ================================================================
# 主入口
# ================================================================

def generate_ai_report(
    df: pd.DataFrame,
    schema_df: pd.DataFrame,
    config: Dict[str, Any],
    analysis_results: Dict[str, Any],
    quality: Optional[Dict[str, Any]],
    provider_config: Dict[str, Any] = None,  # type: ignore[assignment]
    api_key: str = "",
    model: str = "",
    temperature: float = 0.3,
    max_tokens: int = 4096,
    provider_key: str = "",
    chat_path: str = "/chat/completions",
    extra_headers: Optional[Dict[str, str]] = None,
    # ── 现有：报告生成参数 ──
    report_structure: str = "通用调研报告",
    report_style: str = "学术报告风",
    report_length: str = "标准版",
    html_theme: str = "简洁课程作业风",
    custom_options: Optional[Dict[str, Any]] = None,
    # ── 新增：文献综述参数 ──
    literature_config: Optional[Dict[str, Any]] = None,
    # ── 新增：研究背景参数 ──
    background_config: Optional[Dict[str, Any]] = None,
    # ── P1.5：统一配置对象 ──
    llm_config: Any = None,             # Optional[LLMConfig]
    report_config: Any = None,          # Optional[ReportConfig]
) -> Dict[str, Any]:
    """生成 AI 驱动的数据分析报告。

    Args:
        df: 原始数据
        schema_df: 变量类型推断结果
        config: 分析配置（含 report_title, target_variable, group_variables, explanatory_variables 等）
        analysis_results: 统计结果（来自 run_full_analysis）
        quality: 数据质量报告
        provider_config: LLM 厂商配置
        api_key: API Key
        model: 模型名
        temperature: 温度参数（建议 0.2-0.5）
        max_tokens: 最大输出 token
        provider_key: 厂商标识
        chat_path: 自定义 API 路径（仅 custom_openai_compatible）
        extra_headers: 额外请求头（仅 custom_openai_compatible）
        report_structure: 报告结构类型
            - "通用调研报告"
            - "学术论文式报告"
            - "政务决策报告"
            - "商业分析报告"
            - "课程作业报告"
        report_style: 写作语言风格
            - "课程作业风"
            - "政务汇报风"
            - "学术报告风"
            - "商业分析风"
        report_length: 报告篇幅
            - "简短版"
            - "标准版"
            - "详细版"
        html_theme: HTML 导出主题
            - "学术论文白底风"
            - "政务蓝白汇报风"
            - "现代数据看板风"
            - "简洁课程作业风"
            - "商业咨询报告风"
        custom_options: 自定义选项（预留扩展）
        literature_config: 可选。文献综述配置字典：
            - enabled: bool（是否启用文献综述）
            - keywords: str（研究关键词，用于学术数据库检索）
            - max_sources: int（最大文献来源数，5-50）
            - year_range: str（"不限" | "近5年" | "近10年" | "近20年"）
            仅在 report_structure == "学术论文式报告" 时生效。
        background_config: 可选。研究背景配置字典：
            - enabled: bool（是否启用）
            - source_path: str（JSON 调研结果目录路径，或 Markdown 文件路径）
            适用于学术论文式报告和政务决策报告。

    Returns:
        {
            "success": bool,
            "markdown_report": str,    # Markdown 格式报告
            "html_report": str,        # HTML 格式报告
            "docx_report": bytes,      # DOCX 格式报告
            "llm_response": dict,      # LLM 原始响应
            "error": str,              # 失败时
            "warnings": list[str],     # 生成过程中的提示
        }
    """
    result = {
        "success": False,
        "markdown_report": "",
        "html_report": "",
        "docx_report": b"",
        "llm_response": None,
        "error": "",
        "warnings": [],
    }

    # ── P1.5: 统一配置对象优先 → 回退到旧参数 ──
    from src.config_models import LLMConfig, ReportConfig

    if llm_config is not None:
        _llm = llm_config
    else:
        _llm = LLMConfig.from_legacy_kwargs(
            provider_config=provider_config,
            api_key=api_key,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            provider_key=provider_key,
            chat_path=chat_path,
            extra_headers=extra_headers,
        )

    if report_config is not None:
        _rpt = report_config
    else:
        _rpt = ReportConfig.from_legacy_kwargs(
            report_structure=report_structure,
            report_style=report_style,
            report_length=report_length,
            html_theme=html_theme,
            literature_config=literature_config,
            background_config=background_config,
        )

    # ── 预检查 ──
    target = config.get("target_variable", "")
    if not target:
        result["warnings"].append(
            "未指定核心结果变量（target_variable），将生成探索性分析报告。"
        )

    # ---- Step 1: 打包分析结果 ----
    try:
        payload = build_analysis_payload(
            df=df,
            schema_df=schema_df,
            config=config,
            analysis_results=analysis_results,
            quality=quality,
            chart_summaries=None,
            selected_sheet="",
            file_type="",
        )
        # ── v0.1.0 Phase 3.5: 隐私过滤 — 移除不应发送给 LLM 的变量 ──
        from src.analysis_packager import filter_payload_for_ai
        payload = filter_payload_for_ai(payload, schema_df)
    except Exception as e:
        result["error"] = f"打包分析结果失败: {str(e)}"
        return result

    # 检查 payload 中是否有回归和显著性检验
    from src.payload_inspector import (
        payload_has_regression,
        payload_has_significance,
    )
    has_regression = payload_has_regression(payload)
    has_significance = payload_has_significance(payload)

    if not has_regression:
        result["warnings"].append(
            "当前 analysis_payload 中没有回归分析结果，AI 不会撰写回归分析相关章节。"
        )
    if not has_significance:
        result["warnings"].append(
            "当前 analysis_payload 中没有显著性检验结果，AI 不会编造 p 值或显著性判断。"
        )

    # 如果选了学术论文式报告但没有显著性/回归，附加提示
    if _rpt.report_structure == "学术论文式报告" and (not has_regression or not has_significance):
        result["warnings"].append(
            "当前数据仍可生成论文式分析报告，但实证结果部分将以已有描述统计和探索性分析为主。"
        )

    # 检查 payload 大小
    payload_json = to_json_payload(payload)
    payload_bytes = len(payload_json.encode("utf-8"))
    if payload_bytes > 400 * 1024:
        result["warnings"].append(
            f"Analysis payload 较大（{payload_bytes / 1024:.1f} KB），"
            "AI 处理时间可能较长。建议考虑压缩或精简变量。"
        )

    # ---- Step 1.4: 研究背景材料读取（如果启用）----
    background_context: Optional[str] = None

    background_config = _rpt.to_background_config()
    literature_config = _rpt.to_literature_config()

    if background_config and background_config.get("enabled"):
        source_path = background_config.get("source_path", "")
        if source_path:
            try:
                from src.background_research import build_background_context

                background_context = build_background_context(source_path)
                if background_context:
                    result["warnings"].append(
                        f"已从 {source_path} 读取研究背景材料"
                        f"（{len(background_context)} 字符），将注入到报告的研究背景章节。"
                    )
                else:
                    result["warnings"].append(
                        f"背景材料路径 {source_path} 存在但未能提取到有效内容。"
                    )
            except ImportError:
                result["warnings"].append("背景研究模块未找到，已跳过。")
            except Exception as e:
                logger.warning(f"Background research read failed: {e}")
                result["warnings"].append(f"读取研究背景材料失败: {e}")
        else:
            result["warnings"].append("已启用研究背景但未指定来源路径。")

    # ---- Step 1.5: 文献综述检索与合成（如果启用）----
    literature_review_content: Optional[str] = None
    literature_papers: List[Dict[str, Any]] = []

    if literature_config and literature_config.get("enabled"):
        from src.report_options import is_structure_supports_literature
        if not is_structure_supports_literature(_rpt.report_structure):
            result["warnings"].append(
                "文献综述仅在「学术论文式报告」结构下生效。已跳过文献检索。"
            )
        elif not (literature_config.get("keywords") or "").strip():
            result["warnings"].append("未输入研究关键词，跳过文献检索。")
        else:
            try:
                from src.literature_review import generate_literature_review

                # 构建调查上下文
                survey_context = {
                    "report_title": config.get("report_title", ""),
                    "research_subject": config.get("research_subject", ""),
                    "target_variable": config.get("target_variable", ""),
                    "variable_descriptions": build_variable_summary(schema_df, config),
                    "key_findings_summary": build_key_findings_summary(payload),
                }

                lit_result = generate_literature_review(
                    literature_config=literature_config,
                    survey_context=survey_context,
                    **_llm.to_lit_review_kwargs(),
                )

                if lit_result.get("success"):
                    literature_review_content = lit_result["literature_review_text"]
                    literature_papers = lit_result.get("papers_found", [])
                    if lit_result.get("warnings"):
                        result["warnings"].extend(lit_result["warnings"])
                else:
                    result["warnings"].append(
                        f"文献综述生成失败: {lit_result.get('error', '未知错误')}。"
                        "将使用不含文献综述的通用报告结构。"
                    )
            except ImportError:
                result["warnings"].append(
                    "文献综述模块未找到，已跳过文献检索。"
                )
            except Exception as e:
                logger.warning(f"Literature review failed: {e}")
                result["warnings"].append(
                    f"文献综述生成失败: {e}。将使用不含文献综述的通用报告结构。"
                )

    # ---- Step 2: 构建提示词 ----
    try:
        system_prompt, user_prompt = build_ai_report_prompt(
            analysis_payload=payload,
            report_structure=_rpt.report_structure,
            report_style=_rpt.report_style,
            report_length=_rpt.report_length,
            literature_review_content=literature_review_content,
            background_context=background_context,
        )
    except Exception as e:
        result["error"] = f"构建 AI 提示词失败: {str(e)}"
        return result

    # ---- Step 3: 调用 LLM ----
    llm_result = call_llm(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        llm_config=_llm,
    )

    result["llm_response"] = llm_result

    if not llm_result.get("success"):
        result["error"] = f"LLM 调用失败: {llm_result.get('error', '未知错误')}"
        return result

    # ---- Step 3.3: 提取 AI 返回内容（处理 JSON 包裹、转义等）----
    from src.report_rendering import extract_report_content, sanitize_markdown_text
    raw_content = llm_result.get("content", "")
    markdown_text = extract_report_content(raw_content)
    if not markdown_text.strip():
        result["error"] = "LLM 返回了空内容（或仅有 JSON 包裹无实际报告）。"
        return result
    markdown_text = sanitize_markdown_text(markdown_text)

    # ---- Step 3.5: 变量名后处理（自动替换泄露的英文列名）----
    markdown_text = _sanitize_variable_names(markdown_text, schema_df)

    # ---- Step 3.6: 追加参考文献附录（如果使用了文献综述）----
    if literature_papers:
        refs_text = "\n\n---\n\n## 参考文献\n\n"
        for i, paper in enumerate(literature_papers, 1):
            apa = paper.get("apa_citation", paper.get("title", f"Unknown Paper {i}"))
            refs_text += f"[{i}] {apa}\n\n"
        markdown_text += refs_text

    result["markdown_report"] = markdown_text

    # ---- Step 4: 转换为 HTML / DOCX (v0.1.0 Phase 4: 统一渲染管线) ----
    from src.report_rendering import render_markdown_to_html, render_markdown_to_docx
    report_title = config.get("report_title", "数据分析报告")

    try:
        result["html_report"] = render_markdown_to_html(
            markdown_text=markdown_text,
            html_theme=_rpt.html_theme,
            report_title=report_title,
        )
    except Exception as e:
        logger.warning(f"HTML 转换警告: {e}")

    try:
        result["docx_report"] = render_markdown_to_docx(markdown_text, config)
    except Exception as e:
        logger.warning(f"DOCX 转换警告: {e}")

    result["success"] = True

    # ---- Step 5: 保存到 outputs/ ----
    _save_reports(result, config)

    return result


# ================================================================
# 变量名后处理（第三层防护：自动替换泄露的英文列名）
# ================================================================

def _sanitize_variable_names(markdown_text: str, schema_df: pd.DataFrame) -> str:
    """将报告中泄露的英文列名自动替换为中文显示名。

    三层策略：
      1. 匹配 "EnglishName（中文名）" 模式 → 替换为 "中文名"
      2. 构建 col→display_name 映射表，全词替换残留英文列名
      3. 保护代码块（```...``` 和 `...`）内的内容不被替换

    Args:
        markdown_text: AI 生成的 Markdown 报告
        schema_df: 变量类型推断结果（含 column 和 display_name）

    Returns:
        替换后的 Markdown 文本
    """
    if not markdown_text or schema_df is None or schema_df.empty:
        return markdown_text

    # ── 构建所有列名的集合（用于模式匹配）──
    all_cols: set = set()
    name_map: Dict[str, str] = {}
    for _, row in schema_df.iterrows():
        col = str(row.get("column", ""))
        display = str(row.get("display_name", ""))
        if col and len(col) >= 3:
            all_cols.add(col)
            if display and display != col and display != "nan":
                name_map[col] = display
            else:
                name_map[col] = display  # 保留原值，后续用于检测

    if not all_cols:
        return markdown_text

    # ── 保护代码块 ──
    code_blocks: List[str] = []
    placeholder = "___CB{}___"

    def _save_code(m: re.Match) -> str:
        idx = len(code_blocks)
        code_blocks.append(m.group(0))
        return placeholder.format(idx)

    text = re.sub(r'```[^`]*```', _save_code, markdown_text, flags=re.DOTALL)
    text = re.sub(r'``[^`]+``', _save_code, text)
    text = re.sub(r'`[^`]+`', _save_code, text)

    # ── 策略1: 匹配 "EnglishName（中文名）" → "中文名" ──
    # 遍历已知列名，匹配 "col（任意中文内容）" 模式
    for col in sorted(all_cols, key=len, reverse=True):
        # 匹配: col + （ + 非英文内容 + ）
        pattern1 = re.compile(
            r'\b' + re.escape(col) + r'\s*[（(]\s*([^）)]+)\s*[）)]'
        )
        def _replace_mixed(m: re.Match, c=col) -> str:
            inner = m.group(1).strip()
            # 如果括号内是中文（或中英混合但不全是英文变量名），使用括号内容
            if inner and not re.match(r'^[a-z_][a-z0-9_]*$', inner, re.IGNORECASE):
                return inner
            # 括号内也是英文，回到策略2
            return m.group(0)
        text = pattern1.sub(_replace_mixed, text)

    # ── 策略2: 全词替换残留英文列名 ──
    # 只替换 display_name 不同于 col 的情况
    for col in sorted(all_cols, key=len, reverse=True):
        display = name_map.get(col, col)
        if display and display != col and display != "nan":
            pattern = re.compile(r'\b' + re.escape(col) + r'\b')
            text = pattern.sub(display, text)

    # ── 策略3: 对于 display_name==col 的列（无中文映射），删除引导英文名 ──
    # 匹配 "英文名"（中文解释）模式中的英文部分
    for col in sorted(all_cols, key=len, reverse=True):
        display = name_map.get(col, col)
        if display == col or display == "nan":
            # 该列无中文映射，尝试匹配 "col"（...中文...）模式
            p = re.compile(
                r'"' + re.escape(col) + r'"\s*[（(]([^）)]+)[）)]'
            )
            def _clean_quoted(m: re.Match) -> str:
                return m.group(1)  # 只保留中文部分
            text = p.sub(_clean_quoted, text)

    # ── 还原代码块 ──
    for i, block in enumerate(code_blocks):
        text = text.replace(placeholder.format(i), block)

    return text


# ================================================================
# DOCX 生成
# ================================================================

def _markdown_to_docx(markdown_text: str, config: Dict[str, Any]) -> bytes:
    """将 Markdown 报告转为 DOCX 格式。

    Word 报告要求：
      - 标题居中
      - 一级标题加粗
      - 二级标题适度加粗
      - 正文使用常规段落
      - 摘要和关键词部分正常显示
    """
    # 预处理：转义中文星号，防止误渲染为 Markdown 加粗/斜体
    from src.utils import escape_chinese_asterisks
    markdown_text = escape_chinese_asterisks(markdown_text)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx 未安装，无法生成 DOCX。")
        return b""

    title = config.get("report_title", "数据分析报告")
    doc = Document()

    # 标题
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.split("\n")

    for line in lines:
        # 跳过代码块标记
        if line.startswith("```"):
            continue

        # 标题
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.strip() == "":
            continue
        else:
            # 检查是否为表格行
            if "|" in line and line.strip().startswith("|"):
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                # 跳过分隔行
                if all(re.match(r'^:?-{2,}:?$', c) for c in cells):
                    continue
                # 简化表格处理：用制表符分隔的段落
                p = doc.add_paragraph()
                p.add_run(" | ".join(cells))
                continue

            # 处理加粗
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.+?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("本报告由 AI 辅助生成 · 统计结果由程序计算 · 文字由大语言模型撰写")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    # 保存到 bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ================================================================
# 内部：文件保存
# ================================================================

def _save_reports(result: Dict[str, Any], config: Dict[str, Any]) -> None:
    """将生成的报告保存到 outputs/ 目录。"""
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        title_slug = config.get("report_title", "ai_report").replace(" ", "_")[:50]

        # Markdown
        if result.get("markdown_report"):
            md_path = OUTPUT_DIR / f"{title_slug}.md"
            md_path.write_text(result["markdown_report"], encoding="utf-8")

        # HTML
        if result.get("html_report"):
            html_path = OUTPUT_DIR / f"{title_slug}.html"
            html_path.write_text(result["html_report"], encoding="utf-8")

        # DOCX
        if result.get("docx_report"):
            docx_path = OUTPUT_DIR / f"{title_slug}.docx"
            docx_path.write_bytes(result["docx_report"])
    except Exception as e:
        logger.warning(f"保存报告文件失败: {e}")
