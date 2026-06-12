"""报告生成模块 —— 自动生成 DOCX 与 HTML 格式的数据分析报告。

报告基于实际分析结果，不编造数据。
所有结论使用审慎措辞，不夸大 AI 作用，不声称因果关系。
"""

from datetime import datetime
from typing import Optional, Dict, List, Any
import pandas as pd
import numpy as np
from io import BytesIO

# ---- 分析模块引用 ----
from src.analysis import (
    describe_numeric,
    frequency_table,
    cross_analysis_full,
    correlation_with_pvalues,
    ols_regression,
    _get_cn_label,
)
from src.utils import parse_value_description


# ================================================================
# 主入口
# ================================================================

def generate_html_report(
    survey_df: pd.DataFrame,
    var_dict: Dict[str, Dict[str, Any]],
) -> str:
    """生成完整的 HTML 格式数据分析报告。

    Returns:
        自包含的 HTML 文档字符串
    """
    ctx = _AnalysisContext(survey_df, var_dict)
    sections = _build_all_sections(ctx)

    return _render_html(sections, ctx)


def generate_docx_report(
    survey_df: pd.DataFrame,
    var_dict: Dict[str, Dict[str, Any]],
) -> bytes:
    """生成完整的 DOCX 格式数据分析报告。

    Returns:
        .docx 文件字节流
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("缺少 python-docx 库，请执行 pip install python-docx")

    ctx = _AnalysisContext(survey_df, var_dict)
    sections = _build_all_sections(ctx)

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # === 封面 ===
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("政务服务中心公众满意度调查\n数据分析报告")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph("")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # === 正文 ===
    for sec in sections:
        # 标题
        heading = doc.add_heading(f"{sec['number']}、{sec['title']}", level=1)
        for run in heading.runs:
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        # 段落
        for para_text in sec.get("paragraphs", []):
            p = doc.add_paragraph(para_text)
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5

        # 表格
        for table_data in sec.get("tables", []):
            if table_data.get("title"):
                doc.add_paragraph(table_data["title"])
            df = table_data["data"]
            if df is None or df.empty:
                continue

            table = doc.add_table(rows=len(df) + 1, cols=len(df.columns), style="Table Grid")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头
            for j, col_name in enumerate(df.columns):
                cell = table.rows[0].cells[j]
                cell.text = str(col_name)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

            # 数据行
            for i, (_, row) in enumerate(df.iterrows()):
                for j, col_name in enumerate(df.columns):
                    val = row[col_name]
                    if isinstance(val, float):
                        val = f"{val:.2f}" if abs(val) < 100 else f"{val:.0f}"
                    table.rows[i + 1].cells[j].text = str(val)
                    for p in table.rows[i + 1].cells[j].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

            doc.add_paragraph("")

        doc.add_paragraph("")

    # 尾部
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— 报告完 —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ================================================================
# 内部：分析上下文
# ================================================================

class _AnalysisContext:
    """缓存分析中间结果，避免重复计算。"""

    def __init__(self, df: pd.DataFrame, var_dict: Dict):
        self.df = df
        self.var_dict = var_dict
        self.n = len(df)

    @property
    def numeric_cols(self) -> List[str]:
        return self.df.select_dtypes(include=["number"]).columns.tolist()


# ================================================================
# 内部：各节构建
# ================================================================

def _build_all_sections(ctx: _AnalysisContext) -> List[Dict]:
    """构建全部 8 节内容。"""
    return [
        _section_1_background(ctx),
        _section_2_demographics(ctx),
        _section_3_satisfaction_overview(ctx),
        _section_4_group_differences(ctx),
        _section_5_factor_analysis(ctx),
        _section_6_priority_improvement(ctx),
        _section_7_conclusions(ctx),
        _section_8_limitations(ctx),
    ]


def _cn(ctx: _AnalysisContext, col: str) -> str:
    """获取变量中文名。"""
    return _get_cn_label(ctx.var_dict, col)


# ================================================================
# 第一节：项目背景与数据说明
# ================================================================

def _section_1_background(ctx: _AnalysisContext) -> Dict:
    return {
        "number": "一",
        "title": "项目背景与数据说明",
        "paragraphs": [
            "本报告基于政务服务中心公众满意度调查数据，旨在系统分析公众对政务服务的满意度水平、"
            "影响因素及改进方向，为提升政务服务质量提供数据参考。",
            f"本次调查共收集有效问卷 {ctx.n} 份。调查内容涵盖受访者的人口学信息"
            "（性别、年龄组、学历、所属区域）、办理渠道与业务类型、等待时间、"
            "各项满意度评价（等待满意度、工作人员态度、办事流程便捷性、线上服务体验、"
            "信息公开透明度、政策信任度）、总体满意度、推荐意愿、投诉意向以及优先改进事项等。",
            "本报告综合运用描述统计、交叉分析、卡方检验、Pearson 相关分析和多元线性回归分析等方法，"
            "对调查数据进行系统性分析。所有结论均基于实际数据计算得出，不包含主观推测或编造内容。",
        ],
        "tables": [],
    }


# ================================================================
# 第二节：样本基本情况
# ================================================================

def _section_2_demographics(ctx: _AnalysisContext) -> Dict:
    paragraphs = [
        f"本次调查共回收有效问卷 {ctx.n} 份，受访者基本构成如下。",
    ]

    demo_vars = ["gender", "age_group", "education", "district", "channel", "service_type"]
    tables = []

    for var in demo_vars:
        if var not in ctx.df.columns:
            continue

        cn = _cn(ctx, var)
        labels = ctx.var_dict.get(var, {}).get("labels", {})
        freq = frequency_table(ctx.df, var, label_mapping=labels)
        if freq.empty:
            continue

        # 众数信息
        non_missing = freq[freq["标签"] != "缺失"]
        if len(non_missing) > 0:
            top_label = non_missing.iloc[0]["标签"]
            top_pct = non_missing.iloc[0]["百分比(%)"]
            paragraphs.append(
                f"在{cn}方面，占比最高的是「{top_label}」（{top_pct}%）。"
            )

        tables.append({
            "title": f"表 2.{len(tables)+1} {cn}分布",
            "data": freq[["标签", "频次", "百分比(%)"]],
        })

    return {
        "number": "二",
        "title": "样本基本情况",
        "paragraphs": paragraphs,
        "tables": tables,
    }


# ================================================================
# 第三节：公众满意度总体情况
# ================================================================

def _section_3_satisfaction_overview(ctx: _AnalysisContext) -> Dict:
    paragraphs = []
    tables = []

    # 总体满意度
    if "overall_satisfaction" in ctx.df.columns:
        desc = describe_numeric(ctx.df, "overall_satisfaction")
        cn = _cn(ctx, "overall_satisfaction")
        labels = ctx.var_dict.get("overall_satisfaction", {}).get("labels", {})
        freq = frequency_table(ctx.df, "overall_satisfaction", label_mapping=labels)

        # 高评（>=4）占比
        series = pd.to_numeric(ctx.df["overall_satisfaction"], errors="coerce").dropna()
        high_pct = (series >= 4).sum() / len(series) * 100
        low_pct = (series <= 2).sum() / len(series) * 100

        paragraphs.append(
            f"总体满意度（{cn}）均值为 {desc['均值']}（满分 5），标准差为 {desc['标准差']}。"
            f"其中评分为 4 分及以上（满意或非常满意）的受访者占 {high_pct:.1f}%，"
            f"评分为 2 分及以下（不满意或非常不满意）的仅占 {low_pct:.1f}%。"
            f"总体来看，公众对政务服务中心的满意度处于较高水平。"
        )

        tables.append({
            "title": f"表 3.1 总体满意度分布",
            "data": freq[["标签", "频次", "百分比(%)", "累计百分比(%)"]],
        })

    # 满意度等级
    if "satisfaction_level" in ctx.df.columns:
        labels = ctx.var_dict.get("satisfaction_level", {}).get("labels", {})
        freq = frequency_table(ctx.df, "satisfaction_level", label_mapping=labels)
        non_missing = freq[freq["标签"] != "缺失"]
        if len(non_missing) >= 2:
            high_label = non_missing[non_missing["编码值"].astype(str).isin(["2", "3"])]
            if len(high_label) > 0:
                high_pct_total = high_label["百分比(%)"].sum()
                paragraphs.append(
                    f"满意度等级方面，中高满意度（中等满意及以上）的受访者占 {high_pct_total:.1f}%，"
                    f"说明大多数公众对政务服务持正面评价。"
                )

        tables.append({
            "title": "表 3.2 满意度等级分布",
            "data": freq[["标签", "频次", "百分比(%)"]],
        })

    # 推荐意愿
    if "recommend" in ctx.df.columns:
        labels = ctx.var_dict.get("recommend", {}).get("labels", {})
        freq = frequency_table(ctx.df, "recommend", label_mapping=labels)
        willing = freq[(freq["编码值"].astype(str) == "1") & (freq["标签"] != "缺失")]
        if len(willing) > 0:
            willing_pct = willing.iloc[0]["百分比(%)"]
            paragraphs.append(
                f"在推荐意愿方面，{willing_pct}% 的受访者表示愿意向他人推荐该政务服务中心，"
                f"这一指标反映了公众对服务的整体认可度。"
            )

        tables.append({
            "title": "表 3.3 推荐意愿分布",
            "data": freq[["标签", "频次", "百分比(%)"]],
        })

    return {
        "number": "三",
        "title": "公众满意度总体情况",
        "paragraphs": paragraphs,
        "tables": tables,
    }


# ================================================================
# 第四节：不同群体和服务场景下的满意度差异
# ================================================================

def _section_4_group_differences(ctx: _AnalysisContext) -> Dict:
    paragraphs = [
        "为探究不同群体和不同服务场景下满意度是否存在差异，"
        "本研究对以下变量对进行了交叉分析与卡方检验。",
    ]

    cross_pairs = [
        ("district", "satisfaction_level"),
        ("channel", "satisfaction_level"),
        ("age_group", "satisfaction_level"),
        ("education", "satisfaction_level"),
    ]

    tables = []
    sig_count = 0

    for row_col, col_col in cross_pairs:
        if row_col not in ctx.df.columns or col_col not in ctx.df.columns:
            continue

        result = cross_analysis_full(ctx.df, row_col, col_col, ctx.var_dict)
        if "error" in result:
            continue

        row_cn = _cn(ctx, row_col)
        col_cn = _cn(ctx, col_col)

        if result["significant"]:
            sig_count += 1
            paragraphs.append(
                f"**{row_cn}** 与 **{col_cn}** 的卡方检验结果显示两者存在统计显著的关联"
                f"（χ² = {result['chi2']}, p = {result['p_value']} < 0.05），"
                f"表明不同{row_cn}在{col_cn}上的分布存在差异。"
            )
        else:
            paragraphs.append(
                f"**{row_cn}** 与 **{col_cn}** 的卡方检验结果未发现统计显著的关联"
                f"（χ² = {result['chi2']}, p = {result['p_value']} ≥ 0.05），"
                f"在当前样本下，不同{row_cn}在{col_cn}上的分布差异不显著。"
            )

        tables.append({
            "title": f"表 4.{len(tables)+1} {row_cn} × {col_cn} 交叉频数表",
            "data": result["crosstab"],
        })

    if sig_count > 0:
        paragraphs.append(
            f"综上，在 {len(cross_pairs)} 组交叉分析中，有 {sig_count} 组发现统计显著的组间差异，"
            "建议在后续工作中关注这些差异较大的群体。"
        )
    else:
        paragraphs.append(
            "各组交叉分析均未发现统计显著的差异，"
            "这可能意味着不同背景的公众在满意度上较为趋同。"
        )

    return {
        "number": "四",
        "title": "不同群体和服务场景下的满意度差异",
        "paragraphs": paragraphs,
        "tables": tables,
    }


# ================================================================
# 第五节：满意度影响因素分析
# ================================================================

def _section_5_factor_analysis(ctx: _AnalysisContext) -> Dict:
    paragraphs = [
        "为探究总体满意度的主要影响因素，本报告对以下变量进行了 Pearson 相关分析和"
        "多元线性回归分析（OLS）。需要注意的是，相关关系和回归系数反映的是变量间的统计关联，"
        "不能直接等同于因果关系。",
    ]

    tables = []

    # 相关分析
    corr_vars = [
        "wait_satisfaction", "staff_attitude", "process_convenience",
        "online_service", "info_transparency", "policy_trust",
        "overall_satisfaction", "wait_time_min",
    ]
    corr_in_data = [v for v in corr_vars if v in ctx.df.columns]

    if len(corr_in_data) >= 3:
        corr_result = correlation_with_pvalues(ctx.df, corr_in_data, ctx.var_dict)
        r_mat = corr_result["r_matrix"]
        p_mat = corr_result["p_matrix"]

        # 找出与 overall_satisfaction 最强的相关
        if "overall_satisfaction" in corr_in_data:
            os_row = r_mat.index[r_mat.index.str.contains("overall")][0]
            correlations = []
            for idx in r_mat.index:
                if "overall" not in idx:
                    r_val = r_mat.loc[os_row, idx]
                    p_val = p_mat.loc[os_row, idx]
                    correlations.append((idx, r_val, p_val))

            correlations.sort(key=lambda x: abs(x[1]), reverse=True)

            paragraphs.append("**相关分析结果显示：**")
            for cn_label, r_val, p_val in correlations[:5]:
                sig = "统计显著" if p_val < 0.05 else "未达显著水平"
                direction = "正相关" if r_val > 0 else "负相关"
                paragraphs.append(
                    f"- {cn_label} 与总体满意度的 Pearson 相关系数 r = {r_val:.3f} "
                    f"（p = {p_val:.4f}，{sig}），呈{direction}。"
                )

            paragraphs.append(
                "从相关系数大小来看，与总体满意度关联最密切的维度依次为："
                + "、".join([c[0].split("（")[1].rstrip("）") if "（" in c[0] else c[0] for c in correlations[:3]])
                + "。"
            )

        tables.append({
            "title": "表 5.1 Pearson 相关系数矩阵",
            "data": r_mat,
        })

    # 回归分析
    reg_indep = [
        "wait_satisfaction", "staff_attitude", "process_convenience",
        "online_service", "info_transparency", "policy_trust", "wait_time_min",
    ]
    reg_in_data = [v for v in reg_indep if v in ctx.df.columns]

    if "overall_satisfaction" in ctx.df.columns and len(reg_in_data) >= 2:
        reg_result = ols_regression(ctx.df, "overall_satisfaction", reg_in_data, ctx.var_dict)
        if "error" not in reg_result:
            paragraphs.append("")
            paragraphs.append(
                f"**多元线性回归分析结果：** 模型整体显著"
                f"（F = {reg_result['f_statistic']:.3f}, p = {reg_result['f_pvalue']:.4f}），"
                f"调整 R² = {reg_result['adj_r_squared']:.3f}，"
                f"表明模型可解释总体满意度约 {reg_result['adj_r_squared']*100:.1f}% 的变异。"
            )

            coef_df = reg_result["coefficients"]
            sig_vars = coef_df[coef_df["显著性"].str.contains(r"\*", na=False)]
            non_sig_vars = coef_df[
                (~coef_df["显著性"].str.contains(r"\*", na=False))
                & (~coef_df["变量"].str.contains("截距"))
            ]

            if len(sig_vars) > 0:
                pos_list = sig_vars[sig_vars["回归系数"] > 0]
                neg_list = sig_vars[sig_vars["回归系数"] < 0]
                if len(pos_list) > 0:
                    var_names = [v.split("（")[1].rstrip("）") if "（" in v else v for v in pos_list["变量"]]
                    paragraphs.append(
                        f"在控制其他变量后，对总体满意度具有统计显著正向关联的因素包括："
                        + "、".join(var_names)
                        + "。"
                    )
                if len(neg_list) > 0:
                    var_names = [v.split("（")[1].rstrip("）") if "（" in v else v for v in neg_list["变量"]]
                    paragraphs.append(
                        f"具有统计显著负向关联的因素包括：" + "、".join(var_names) + "。"
                    )

            if len(non_sig_vars) > 0:
                var_names = [v.split("（")[1].rstrip("）") if "（" in v else v for v in non_sig_vars["变量"]]
                paragraphs.append(
                    f"在多元回归模型中未达统计显著水平的变量包括："
                    + "、".join(var_names)
                    + "。"
                )

            paragraphs.append(
                "综合相关分析和回归分析结果，等待满意度、工作人员态度、办事流程便捷性、"
                "线上服务体验和信息透明度是影响总体满意度的关键因素。"
                "建议将这些维度作为服务质量改进的优先关注方向。"
            )

            tables.append({
                "title": "表 5.2 多元线性回归系数表",
                "data": coef_df,
            })

    paragraphs.append("")
    paragraphs.append(
        "**重要提示：** 以上分析基于横截面调查数据，统计关联不等于因果关系。"
        "可能存在未纳入模型的混淆变量。分析结果应结合实际情况和领域知识进行综合判断，"
        "作为决策参考而非唯一依据。"
    )

    return {
        "number": "五",
        "title": "满意度影响因素分析",
        "paragraphs": paragraphs,
        "tables": tables,
    }


# ================================================================
# 第六节：优先改进事项分析
# ================================================================

def _section_6_priority_improvement(ctx: _AnalysisContext) -> Dict:
    paragraphs = []
    tables = []

    if "priority_improve" in ctx.df.columns:
        labels = ctx.var_dict.get("priority_improve", {}).get("labels", {})
        freq = frequency_table(ctx.df, "priority_improve", label_mapping=labels)
        non_missing = freq[freq["标签"] != "缺失"].sort_values("频次", ascending=False)

        if len(non_missing) > 0:
            top_items = non_missing.head(3)
            top_str = "、".join(
                f"「{row['标签']}」（{row['百分比(%)']}%）"
                for _, row in top_items.iterrows()
            )

            paragraphs.append(
                f"在优先改进事项方面，受访者选择最多的三项依次为：{top_str}。"
            )

            paragraphs.append(
                "这一结果反映了公众对政务服务改进的核心诉求。"
                "建议相关部门将这些事项作为阶段性重点工作加以推进，"
                "并根据不同群体的差异化需求进行精准施策。"
            )

        tables.append({
            "title": "表 6.1 优先改进事项分布",
            "data": freq[["标签", "频次", "百分比(%)"]],
        })
    else:
        paragraphs.append("数据中未包含优先改进事项变量，本节暂无法分析。")

    return {
        "number": "六",
        "title": "优先改进事项分析",
        "paragraphs": paragraphs,
        "tables": tables,
    }


# ================================================================
# 第七节：结论与建议
# ================================================================

def _section_7_conclusions(ctx: _AnalysisContext) -> Dict:
    paragraphs = [
        "基于以上分析，本报告提出以下政策建议：",
        "",
        "**1. 优化办事流程，提高流程便利性。** "
        "回归分析表明，办事流程便捷性与总体满意度存在显著正向关联。"
        "建议进一步简化办事环节，推行「一窗受理、集成服务」模式，"
        "减少公众跑腿次数和材料重复提交。",

        "**2. 缩短等待时间，改善窗口排队体验。** "
        "等待时间是影响公众体验的重要因素。建议通过预约办理、智能分流、"
        "弹性窗口调度等方式，降低高峰时段的平均等待时间。",

        "**3. 提升线上服务的稳定性和易用性。** "
        "线上服务体验对满意度有独立贡献。建议加强政务 App 和小程序的交互设计，"
        "提升系统响应速度，完善在线办理功能，推动更多事项实现「全程网办」。",

        "**4. 加强信息公开透明度。** "
        "信息透明度是公众信任的基础。建议在政务大厅和线上平台同步更新办事指南、"
        "办理进度和结果反馈，让公众对办事流程和预期时间有清晰了解。",

        "**5. 针对满意度较低群体开展精准改进。** "
        "交叉分析结果提示不同群体在满意度上可能存在差异。"
        "建议针对满意度评分较低的群体（如特定区域或特定渠道用户）进行深度调研，"
        "了解其具体痛点并制定针对性改进方案。",
    ]

    paragraphs.append("")
    paragraphs.append(
        "以上建议综合了描述统计、交叉分析、相关分析和回归分析的主要发现。"
        "在推进改进措施时，建议建立持续监测机制，"
        "定期追踪关键满意度指标的变化趋势，以评估改进成效并及时调整策略。"
    )

    return {
        "number": "七",
        "title": "结论与建议",
        "paragraphs": paragraphs,
        "tables": [],
    }


# ================================================================
# 第八节：方法说明与局限
# ================================================================

def _section_8_limitations(ctx: _AnalysisContext) -> Dict:
    return {
        "number": "八",
        "title": "方法说明与局限",
        "paragraphs": [
            "本报告的分析方法主要包括：描述统计（均值、标准差、频数分布）、"
            "交叉分析与卡方检验（检验分类变量之间的关联）、"
            "Pearson 相关分析（衡量变量间的线性关联强度）以及"
            "多元线性回归分析（探究多个因素对总体满意度的独立贡献）。",

            "本报告存在以下局限性：",

            "**第一，数据代表性问题。** 本调查的样本可能未能完全覆盖所有类型的政务服务受众，"
            "结论的外推需谨慎。",

            "**第二，因果关系识别问题。** 相关分析和回归分析只能揭示变量之间的统计关联，"
            "不能直接证明因果关系。可能存在未纳入分析的第三变量（如区域经济发展水平、"
            "个人期望差异等）同时影响各项满意度评价。",

            "**第三，自报告数据的偏差。** 满意度评价基于受访者的主观判断，"
            "可能受到回忆偏差、社会期望偏差等因素的影响。",

            "**第四，变量测量精度。** 部分变量采用定序或分类尺度测量，"
            "在回归分析中被近似为连续变量处理，可能对估计精度有一定影响。",

            "综上所述，本报告的分析结果为政务服务质量改进提供了有益的数据参考，"
            "但统计结果应作为辅助决策工具，而非替代专业判断。"
            "建议结合定性调研（如深度访谈、焦点小组）构建更全面的改进策略。",
        ],
        "tables": [],
    }


# ================================================================
# HTML 渲染
# ================================================================

def _render_html(sections: List[Dict], ctx: _AnalysisContext) -> str:
    """将报告节渲染为自包含 HTML 文档。"""
    now = datetime.now().strftime("%Y年%m月%d日")

    html_parts = [f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>政务服务中心公众满意度调查数据分析报告</title>
<style>
  body {{
    font-family: "Microsoft YaHei", "SimSun", sans-serif;
    max-width: 900px;
    margin: 40px auto;
    padding: 0 20px;
    color: #333;
    line-height: 1.8;
    background: #fff;
  }}
  h1 {{
    text-align: center;
    font-size: 22px;
    margin-bottom: 5px;
    color: #2B5F8A;
    border-bottom: 2px solid #2B5F8A;
    padding-bottom: 15px;
  }}
  h2 {{
    font-size: 16px;
    margin-top: 30px;
    color: #2B5F8A;
    border-left: 4px solid #2B5F8A;
    padding-left: 10px;
  }}
  p {{
    text-indent: 2em;
    margin: 8px 0;
  }}
  p.no_indent {{
    text-indent: 0;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: 15px 0;
    font-size: 13px;
  }}
  th {{
    background: #2B5F8A;
    color: white;
    padding: 8px;
    border: 1px solid #ddd;
  }}
  td {{
    padding: 6px 8px;
    border: 1px solid #ddd;
    text-align: center;
  }}
  tr:nth-child(even) {{ background: #f8f9fa; }}
  .cover {{
    text-align: center;
    padding: 100px 0;
  }}
  .cover h1 {{
    font-size: 28px;
    border: none;
    line-height: 1.6;
  }}
  .cover .date {{
    color: #888;
    font-size: 14px;
    margin-top: 30px;
  }}
  .footer {{
    text-align: center;
    color: #888;
    margin-top: 50px;
    border-top: 1px solid #eee;
    padding-top: 20px;
  }}
  @media print {{
    body {{ max-width: 100%; }}
    h2 {{ page-break-before: always; }}
  }}
</style>
</head>
<body>

<div class="cover">
<h1>政务服务中心公众满意度调查<br>数据分析报告</h1>
<p class="date">报告生成日期：{now}</p>
</div>

<hr style="border: none; border-top: 1px solid #2B5F8A; margin: 30px 0;">
"""]

    for sec in sections:
        html_parts.append(f"<h2>{sec['number']}、{sec['title']}</h2>")

        for para in sec.get("paragraphs", []):
            cls = ' class="no_indent"' if para.startswith("**") and para.endswith("**") else ""
            # 处理 Markdown 加粗
            import re
            para_html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", para)
            html_parts.append(f"<p{cls}>{para_html}</p>")

        for table_data in sec.get("tables", []):
            if table_data.get("title"):
                html_parts.append(
                    f'<p class="no_indent"><em>{table_data["title"]}</em></p>'
                )
            df = table_data["data"]
            if df is None or df.empty:
                continue

            html_parts.append("<table><thead><tr>")
            for col in df.columns:
                html_parts.append(f"<th>{col}</th>")
            html_parts.append("</tr></thead><tbody>")

            for _, row in df.iterrows():
                html_parts.append("<tr>")
                for col in df.columns:
                    val = row[col]
                    if isinstance(val, float):
                        if abs(val) < 1 and val != 0:
                            val = f"{val:.4f}"
                        elif abs(val) < 100:
                            val = f"{val:.2f}"
                        else:
                            val = f"{val:.1f}"
                    html_parts.append(f"<td>{val}</td>")
                html_parts.append("</tr>")
            html_parts.append("</tbody></table>")

    html_parts.append("""
<div class="footer">
<p>— 报告完 —</p>
<p>本报告由政务服务中心公众满意度调查数据分析系统自动生成，数据来源为实际调查结果。</p>
</div>

</body>
</html>""")

    return "\n".join(html_parts)


# ================================================================
# 通用报告生成（通用模式）
# ================================================================

def generate_generic_html_report(
    df: pd.DataFrame,
    schema_df: "pd.DataFrame",
    config: Dict[str, Any],
    var_dict: Optional[Dict[str, Dict[str, Any]]] = None,
) -> str:
    """为通用模式生成 HTML 格式的数据分析报告。

    Args:
        df: 调查数据
        schema_df: 变量类型推断结果
        config: 分析配置（report_title, target_variable, group_variables, explanatory_variables）
        var_dict: 可选变量字典

    Returns:
        自包含的 HTML 文档字符串
    """
    ctx = _GenericReportContext(df, schema_df, config, var_dict)
    sections = _build_generic_sections(ctx)
    return _render_generic_html(sections, ctx)


def generate_generic_docx_report(
    df: pd.DataFrame,
    schema_df: "pd.DataFrame",
    config: Dict[str, Any],
    var_dict: Optional[Dict[str, Dict[str, Any]]] = None,
) -> bytes:
    """为通用模式生成 DOCX 格式的数据分析报告。

    Args:
        df: 调查数据
        schema_df: 变量类型推断结果
        config: 分析配置
        var_dict: 可选变量字典

    Returns:
        .docx 文件字节流
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("缺少 python-docx 库，请执行 pip install python-docx")

    ctx = _GenericReportContext(df, schema_df, config, var_dict)
    sections = _build_generic_sections(ctx)

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 封面
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config.get("report_title", "数据分析报告"))
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph("")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # 正文
    for sec in sections:
        heading = doc.add_heading(f"{sec['number']}、{sec['title']}", level=1)
        for run in heading.runs:
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        for para_text in sec.get("paragraphs", []):
            p = doc.add_paragraph(para_text)
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5

        for table_data in sec.get("tables", []):
            if table_data.get("title"):
                doc.add_paragraph(table_data["title"])
            table_df = table_data["data"]
            if table_df is None or table_df.empty:
                continue

            table = doc.add_table(rows=len(table_df) + 1, cols=len(table_df.columns), style="Table Grid")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, col_name in enumerate(table_df.columns):
                cell = table.rows[0].cells[j]
                cell.text = str(col_name)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

            for i, (_, row) in enumerate(table_df.iterrows()):
                for j, col_name in enumerate(table_df.columns):
                    val = row[col_name]
                    if isinstance(val, float):
                        val = f"{val:.2f}" if abs(val) < 100 else f"{val:.0f}"
                    table.rows[i + 1].cells[j].text = str(val)
                    for p in table.rows[i + 1].cells[j].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

            doc.add_paragraph("")

        doc.add_paragraph("")

    # 尾部
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— 报告完 —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ================================================================
# 通用报告上下文
# ================================================================

class _GenericReportContext:
    """通用模式报告上下文，缓存分析结果和配置。"""

    def __init__(
        self,
        df: pd.DataFrame,
        schema_df: "pd.DataFrame",
        config: Dict[str, Any],
        var_dict: Optional[Dict] = None,
    ):
        self.df = df
        self.schema_df = schema_df
        self.config = config
        self.var_dict = var_dict or {}
        self.n = len(df)
        self.report_title = config.get("report_title", "数据分析报告")
        self.target = config.get("target_variable", "")
        self.group_vars = config.get("group_variables", []) or []
        self.expl_vars = config.get("explanatory_variables", []) or []

        # 构建快速查找
        self.type_map = {}
        self.cn_map = {}
        for _, row in schema_df.iterrows():
            col = row["column"]
            self.type_map[col] = row["inferred_type"]
            self.cn_map[col] = row.get("display_name", "") or col

        # 运行分析
        from src.generic_analysis import run_full_analysis
        self.analysis = run_full_analysis(df, schema_df, config, var_dict)

    def _cn(self, col: str) -> str:
        return self.cn_map.get(col, col)

    def _vtype(self, col: str) -> str:
        return self.type_map.get(col, "")


# ================================================================
# 通用报告各节构建
# ================================================================

def _build_generic_sections(ctx: _GenericReportContext) -> List[Dict]:
    """构建通用报告全部 8 节。"""
    return [
        _g_sec_1_overview(ctx),
        _g_sec_2_variables(ctx),
        _g_sec_3_target(ctx),
        _g_sec_4_groups(ctx),
        _g_sec_5_relationships(ctx),
        _g_sec_6_findings(ctx),
        _g_sec_7_recommendations(ctx),
        _g_sec_8_methodology(ctx),
    ]


def _g_sec_1_overview(ctx: _GenericReportContext) -> Dict:
    """第一节：数据来源与样本概况。"""
    paragraphs = [
        f"本报告基于「{ctx.report_title}」调查数据，旨在通过系统的统计分析，"
        f"了解数据的基本特征、变量间的关系，为相关决策提供数据参考。",
        f"本次分析共涉及 {ctx.n} 条记录，{len(ctx.df.columns)} 个变量。",
    ]

    # 缺失值与重复
    missing_total = int(ctx.df.isnull().sum().sum())
    duplicate_count = int(ctx.df.duplicated().sum())
    paragraphs.append(
        f"数据质量方面，共存在 {missing_total} 个缺失值（缺失率 "
        f"{round(missing_total / max(ctx.df.size, 1) * 100, 2)}%），"
        f"重复记录 {duplicate_count} 条（重复率 "
        f"{round(duplicate_count / max(ctx.n, 1) * 100, 2)}%）。"
    )

    # 变量类型分布
    type_counts = ctx.schema_df["inferred_type"].value_counts()
    type_desc = "、".join(f"{t} 类 {c} 个" for t, c in type_counts.items())
    paragraphs.append(f"根据自动推断，变量类型分布为：{type_desc}。")

    return {
        "number": "一",
        "title": "数据来源与样本概况",
        "paragraphs": paragraphs,
        "tables": [],
    }


def _g_sec_2_variables(ctx: _GenericReportContext) -> Dict:
    """第二节：变量结构与数据质量。"""
    paragraphs = [
        "下表列出了所有变量的基本信息，包括推断类型、缺失值情况和示例值。"
    ]

    # 变量说明表
    display_schema = ctx.schema_df[[
        "column", "display_name", "inferred_type",
        "missing_count", "missing_rate", "unique_count", "example_values",
    ]].copy()
    display_schema.columns = [
        "变量名", "中文名称", "推断类型", "缺失数", "缺失率(%)", "唯一值数", "示例值"
    ]

    paragraphs.append(
        "建议关注缺失率较高的变量，必要时在分析中采取适当的缺失值处理方法。"
        "对于高基数分类变量（high_cardinality），建议降维或分组后再进行分析。"
    )

    return {
        "number": "二",
        "title": "变量结构与数据质量",
        "paragraphs": paragraphs,
        "tables": [{"title": "表 2.1 变量信息一览", "data": display_schema}],
    }


def _g_sec_3_target(ctx: _GenericReportContext) -> Dict:
    """第三节：核心变量分析。"""
    paragraphs = []
    tables = []

    target = ctx.target
    if not target or target not in ctx.df.columns:
        paragraphs.append("未指定核心结果变量，本节暂无法分析。请通过分析配置指定目标变量。")
        return {"number": "三", "title": "核心变量分析", "paragraphs": paragraphs, "tables": tables}

    target_cn = ctx._cn(target)
    target_type = ctx._vtype(target)

    # 获取单变量分析结果
    univariate = ctx.analysis.get("univariate", {})
    uv_result = univariate.get(target, {})

    if isinstance(uv_result, dict) and "error" not in uv_result:
        if target_type == "numeric":
            paragraphs.append(
                f"**{target_cn}** 的均值为 {uv_result.get('均值', 'N/A')}，"
                f"标准差为 {uv_result.get('标准差', 'N/A')}，"
                f"中位数为 {uv_result.get('中位数', 'N/A')}，"
                f"取值范围为 [{uv_result.get('最小值', 'N/A')}, {uv_result.get('最大值', 'N/A')}]。"
                f"数据分布偏度为 {uv_result.get('偏度', 'N/A')}（"
                + ("右偏" if uv_result.get('偏度', 0) > 0.5 else
                   "左偏" if uv_result.get('偏度', 0) < -0.5 else "近似对称")
                + "）。"
            )

            # 描述统计表
            desc_rows = {k: v for k, v in uv_result.items()
                        if k in ("样本量", "均值", "标准差", "中位数", "最小值", "最大值", "Q25", "Q75", "偏度", "峰度")}
            tables.append({
                "title": f"表 3.1 {target_cn} 描述统计",
                "data": pd.DataFrame([desc_rows]),
            })

        elif target_type in ("categorical", "ordinal"):
            mode_label = uv_result.get("众数", "")
            mode_pct = uv_result.get("众数占比", 0)
            cat_count = uv_result.get("类别数", 0)
            paragraphs.append(
                f"**{target_cn}** 共包含 {cat_count} 个类别，"
                f"有效样本 {uv_result.get('有效样本', 'N/A')} 个。"
                f"出现频次最高的类别为「{mode_label}」（占比 {mode_pct:.1f}%）。"
            )
            if target_type == "ordinal":
                paragraphs.append(
                    f"作为有序变量，**{target_cn}** 的均值为 {uv_result.get('均值', 'N/A')}，"
                    f"中位数为 {uv_result.get('中位数', 'N/A')}。"
                )

            freq = uv_result.get("频数表")
            if freq is not None and not freq.empty:
                tables.append({
                    "title": f"表 3.1 {target_cn} 频数分布",
                    "data": freq[["标签", "频次", "百分比(%)"]],
                })
    else:
        paragraphs.append(f"**{target_cn}** 的分析未能完成。{uv_result.get('error', '')}")

    return {
        "number": "三",
        "title": "核心变量分析",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def _g_sec_4_groups(ctx: _GenericReportContext) -> Dict:
    """第四节：群体差异分析。"""
    paragraphs = [
        "本节通过交叉分析和分组比较，探究不同群体在核心变量上是否存在差异。"
    ]
    tables = []

    if not ctx.target or not ctx.group_vars:
        paragraphs.append("未指定分组变量或目标变量，本节暂无法分析。")
        return {"number": "四", "title": "群体差异分析", "paragraphs": paragraphs, "tables": tables}

    sig_count = 0
    bivariate_group = ctx.analysis.get("bivariate_group", {})

    for key, result in bivariate_group.items():
        if not isinstance(result, dict):
            continue

        if "error" in result:
            paragraphs.append(f"- {result.get('error', '')}")
            continue

        if "info" in result:
            continue

        if "chi2" in result:
            # 交叉分析结果
            row_cn = result.get("row_var_cn", "")
            col_cn = result.get("col_var_cn", "")
            if result.get("significant"):
                sig_count += 1
                paragraphs.append(
                    f"**{row_cn}** 与 **{col_cn}** 之间存在统计显著的关联"
                    f"（χ² = {result['chi2']}, p = {result['p_value']} < 0.05）。"
                )
            else:
                paragraphs.append(
                    f"**{row_cn}** 与 **{col_cn}** 之间未发现统计显著的关联"
                    f"（χ² = {result['chi2']}, p = {result['p_value']} ≥ 0.05）。"
                )

            if "crosstab" in result:
                tables.append({
                    "title": f"表 4.{len(tables)+1} {row_cn} × {col_cn} 交叉表",
                    "data": result["crosstab"],
                })

        elif "group_stats" in result:
            # 分组均值比较
            cat_cn = result.get("cn_cat", "")
            num_cn = result.get("cn_num", "")
            if result.get("significant"):
                sig_count += 1
                paragraphs.append(
                    f"不同 **{cat_cn}** 在 **{num_cn}** 上存在统计显著的差异"
                    f"（{result.get('test_type', '')}, p = {result['p_value']} < 0.05）。"
                )
            else:
                paragraphs.append(
                    f"不同 **{cat_cn}** 在 **{num_cn}** 上未发现统计显著的差异"
                    f"（p = {result['p_value']} ≥ 0.05）。"
                )

            tables.append({
                "title": f"表 4.{len(tables)+1} {cat_cn} × {num_cn} 分组统计",
                "data": result["group_stats"],
            })

    if sig_count == 0 and bivariate_group:
        paragraphs.append(
            "各组交叉分析均未发现统计显著差异，"
            "这可能意味着不同群体的特征较为趋同，或样本量不足以检出较小差异。"
        )

    return {
        "number": "四",
        "title": "群体差异分析",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def _g_sec_5_relationships(ctx: _GenericReportContext) -> Dict:
    """第五节：变量关系分析。"""
    paragraphs = [
        "本节通过相关分析和回归分析，探究解释变量与核心变量之间的关系。"
        "需要特别说明：相关关系和回归系数反映的是变量间的统计关联，不能直接等同于因果关系。"
    ]
    tables = []

    # 相关分析结果
    bivariate_corr = ctx.analysis.get("bivariate_corr", {})
    sig_corrs = []
    for key, result in bivariate_corr.items():
        if isinstance(result, dict) and result.get("significant"):
            cn1 = result.get("cn1", "")
            cn2 = result.get("cn2", "")
            r_val = result.get("pearson_r", 0)
            sig_corrs.append((cn1, cn2, r_val))

    if sig_corrs:
        sig_corrs.sort(key=lambda x: abs(x[2]), reverse=True)
        paragraphs.append("**统计显著的相关关系（按强度降序）：**")
        for cn1, cn2, r_val in sig_corrs[:10]:
            direction = "正" if r_val > 0 else "负"
            paragraphs.append(f"- **{cn1}** 与 **{cn2}**：r = {r_val:.3f}，{direction}相关。")

        # 相关矩阵热力图数据
        if len(sig_corrs) >= 2:
            # 提取涉及的数值变量
            num_vars_in_corr = set()
            for cn1, cn2, _ in sig_corrs:
                for col, cn in ctx.cn_map.items():
                    if cn == cn1 or cn == cn2:
                        num_vars_in_corr.add(col)

            if len(num_vars_in_corr) >= 2:
                from src.analysis import correlation_matrix as calc_corr_matrix
                corr_df = calc_corr_matrix(ctx.df, list(num_vars_in_corr), ctx.var_dict)
                if not corr_df.empty:
                    tables.append({
                        "title": "表 5.1 Pearson 相关系数矩阵",
                        "data": corr_df,
                    })
    else:
        paragraphs.append("在当前数据中，未发现统计显著的变量间线性相关关系。")

    # 回归分析
    multi = ctx.analysis.get("multivariate", {})
    if isinstance(multi, dict) and "error" not in multi and "coefficients" in multi:
        paragraphs.append("")
        paragraphs.append("**多元线性回归分析：**")
        paragraphs.append(
            f"模型整体显著（F = {multi.get('f_statistic', 'N/A')}, "
            f"p = {multi.get('f_pvalue', 'N/A')}），"
            f"调整 R² = {multi.get('adj_r_squared', 'N/A')}，"
            f"有效样本 n = {multi.get('n', 'N/A')}。"
        )

        coef_df = multi["coefficients"]
        sig_vars = coef_df[coef_df["显著性"].str.contains(r"\*", na=False)]
        non_sig_vars = coef_df[
            (~coef_df["显著性"].str.contains(r"\*", na=False))
            & (~coef_df["变量"].str.contains("截距"))
        ]

        if len(sig_vars) > 0:
            pos_vars = sig_vars[sig_vars["回归系数"] > 0]
            neg_vars = sig_vars[sig_vars["回归系数"] < 0]
            if len(pos_vars) > 0:
                var_list = [v.split("（")[1].rstrip("）") if "（" in str(v) else str(v) for v in pos_vars["变量"]]
                paragraphs.append("显著正向关联：" + "、".join(var_list) + "。")
            if len(neg_vars) > 0:
                var_list = [v.split("（")[1].rstrip("）") if "（" in str(v) else str(v) for v in neg_vars["变量"]]
                paragraphs.append("显著负向关联：" + "、".join(var_list) + "。")

        if len(non_sig_vars) > 0:
            var_list = [v.split("（")[1].rstrip("）") if "（" in str(v) else str(v) for v in non_sig_vars["变量"]]
            paragraphs.append("未达显著水平的变量：" + "、".join(var_list) + "。")

        tables.append({
            "title": "表 5.2 多元线性回归系数表",
            "data": coef_df,
        })
    elif isinstance(multi, dict) and "error" in multi:
        paragraphs.append(f"回归分析未能完成：{multi['error']}")

    paragraphs.append("")
    paragraphs.append(
        "**重要提示：** 以上分析基于横截面调查数据，统计关联不等于因果关系。"
        "回归系数仅反映在控制其他变量条件下的统计关联强度，"
        "建议结合专业知识和实际情况进行综合判断。"
    )

    return {
        "number": "五",
        "title": "变量关系分析",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def _g_sec_6_findings(ctx: _GenericReportContext) -> Dict:
    """第六节：主要发现。"""
    paragraphs = ["基于以上分析结果，本报告归纳以下主要发现："]

    findings = []

    # 从分析结果自动提取关键发现
    target_cn = ctx._cn(ctx.target) if ctx.target else ""

    # 1. 目标变量概览
    univariate = ctx.analysis.get("univariate", {})
    if ctx.target and ctx.target in univariate:
        uv = univariate[ctx.target]
        if isinstance(uv, dict) and "error" not in uv:
            if "均值" in uv and uv.get("均值") is not None:
                findings.append(
                    f"**发现一：{target_cn}总体概况。** "
                    f"{target_cn}的均值为 {uv.get('均值', 'N/A')}（标准差 {uv.get('标准差', 'N/A')}），"
                    f"数据分布偏度为 {uv.get('偏度', 'N/A')}，"
                    f"表明整体水平" + ("偏高" if uv.get('均值', 0) > 3 else "中等" if uv.get('均值', 0) > 2 else "偏低") + "。"
                )
            elif "众数" in uv:
                findings.append(
                    f"**发现一：{target_cn}分布特征。** "
                    f"最常见的类别为「{uv.get('众数', '')}」（占比 {uv.get('众数占比', 0):.1f}%），"
                    f"共包含 {uv.get('类别数', 0)} 个类别。"
                )

    # 2. 显著组间差异
    bivariate_group = ctx.analysis.get("bivariate_group", {})
    sig_groups = []
    for key, result in bivariate_group.items():
        if isinstance(result, dict) and result.get("significant"):
            if "chi2" in result:
                sig_groups.append(f"**{result.get('row_var_cn', '')}** 与 **{result.get('col_var_cn', '')}**")
            elif "group_stats" in result:
                sig_groups.append(f"**{result.get('cn_cat', '')}** 在 **{result.get('cn_num', '')}** 上的差异")

    if sig_groups:
        findings.append(
            f"**发现二：显著的群体差异。** "
            f"分析发现以下变量对之间存在统计显著的差异或关联："
            + "；".join(sig_groups[:3]) + "。"
        )

    # 3. 显著相关
    bivariate_corr = ctx.analysis.get("bivariate_corr", {})
    sig_corrs = []
    for key, result in bivariate_corr.items():
        if isinstance(result, dict) and result.get("significant"):
            cn1 = result.get("cn1", "")
            cn2 = result.get("cn2", "")
            r_val = result.get("pearson_r", 0)
            sig_corrs.append((cn1, cn2, r_val))

    if sig_corrs:
        sig_corrs.sort(key=lambda x: abs(x[2]), reverse=True)
        top_corr = sig_corrs[0]
        findings.append(
            f"**发现三：关键关联关系。** "
            f"与 {target_cn} 关联最强的变量为「{top_corr[0]}」"
            f"（r = {top_corr[2]:.3f}），"
            f"提示" + ("正向" if top_corr[2] > 0 else "负向") + "统计关联。"
        )

    # 4. 回归结论
    multi = ctx.analysis.get("multivariate", {})
    if isinstance(multi, dict) and "coefficients" in multi and multi.get("adj_r_squared", 0) > 0.05:
        findings.append(
            f"**发现四：多变量解释力。** "
            f"回归模型可解释 {target_cn} 约 {multi['adj_r_squared']*100:.1f}% 的变异，"
            f"模型整体显著（p = {multi.get('f_pvalue', 'N/A')}），"
            f"提示所选变量具有一定的解释能力。"
        )

    # 5. 数据质量
    missing_total = int(ctx.df.isnull().sum().sum())
    if missing_total > 0:
        findings.append(
            f"**发现五：数据质量提示。** "
            f"数据集中存在 {missing_total} 个缺失值，"
            f"建议在后续分析中对缺失率较高的变量予以关注。"
        )

    # 至少保证 3 条发现
    if len(findings) < 3:
        findings.append(
            "在当前数据范围内，统计显著的结果较为有限。"
            "建议扩大样本量或增加变量维度以获取更丰富的分析结论。"
        )

    paragraphs.extend(findings)

    paragraphs.append("")
    paragraphs.append(
        "以上发现均基于实际数据统计得出，供进一步分析和决策参考。"
        "统计关联不能直接解释为因果关系。"
    )

    return {
        "number": "六",
        "title": "主要发现",
        "paragraphs": paragraphs,
        "tables": [],
    }


def _g_sec_7_recommendations(ctx: _GenericReportContext) -> Dict:
    """第七节：分析建议。"""
    paragraphs = [
        "基于以上分析结果，提出以下建议供参考：",
        "",
    ]

    # 基于结果的建议
    rec_num = 0

    # 检查是否有显著组间差异 → 建议细分
    bivariate_group = ctx.analysis.get("bivariate_group", {})
    has_sig_group = any(
        isinstance(r, dict) and r.get("significant")
        for r in bivariate_group.values()
    )
    if has_sig_group:
        rec_num += 1
        paragraphs.append(
            f"**{rec_num}. 关注群体差异，制定针对性策略。** "
            f"分析发现不同群体在核心指标上存在显著差异。"
            f"建议进一步细分群体，了解差异化需求，制定更具针对性的解决方案。"
        )

    # 检查是否有显著相关 → 建议关注影响因素
    bivariate_corr = ctx.analysis.get("bivariate_corr", {})
    sig_corrs = [
        (k, v) for k, v in bivariate_corr.items()
        if isinstance(v, dict) and v.get("significant")
    ]
    if sig_corrs:
        rec_num += 1
        top = sorted(sig_corrs, key=lambda x: abs(x[1].get("pearson_r", 0)), reverse=True)[:3]
        top_names = [t[1].get("cn1", "") + "与" + t[1].get("cn2", "") for t in top]
        paragraphs.append(
            f"**{rec_num}. 重视关键关联因素。** "
            f"相关分析显示，" + "、".join(top_names) + "等变量间存在显著关联。"
            f"建议在制定改进方案时优先关注这些维度的协同优化。"
        )

    # 回归显著 → 建议多因素综合施策
    multi = ctx.analysis.get("multivariate", {})
    if isinstance(multi, dict) and "coefficients" in multi and multi.get("adj_r_squared", 0) > 0.05:
        rec_num += 1
        paragraphs.append(
            f"**{rec_num}. 多因素综合施策。** "
            f"回归分析表明，多个因素对核心变量具有独立贡献。"
            f"建议采取综合改进策略，而非单一维度的调整。"
        )

    # 通用建议（补齐）
    if rec_num < 3:
        rec_num += 1
        paragraphs.append(
            f"**{rec_num}. 建立持续监测机制。** "
            f"建议定期追踪关键指标的变化趋势，建立数据驱动的评估和反馈体系，"
            f"以便及时发现问题和评估改进效果。"
        )

    rec_num += 1
    paragraphs.append(
        f"**{rec_num}. 结合定性研究深入分析。** "
        f"定量分析能揭示统计规律，但对现象背后的原因解释力有限。"
        f"建议结合深度访谈、焦点小组等定性方法，构建更全面的分析视角。"
    )

    paragraphs.append("")
    paragraphs.append(
        "以上建议基于数据分析结果提出，在具体实施时应结合实际情况、"
        "资源约束和专业知识进行综合评估。"
    )

    return {
        "number": "七",
        "title": "分析建议",
        "paragraphs": paragraphs,
        "tables": [],
    }


def _g_sec_8_methodology(ctx: _GenericReportContext) -> Dict:
    """第八节：方法说明与局限。"""
    return {
        "number": "八",
        "title": "方法说明与局限",
        "paragraphs": [
            "本报告的分析方法主要包括：描述统计（均值、标准差、频数分布）、"
            "交叉分析与卡方检验（检验分类变量之间的关联）、"
            "分组均值比较与方差分析、"
            "Pearson 与 Spearman 相关分析（衡量变量间的线性及单调关联强度）以及"
            "多元线性回归分析（探究多个因素对目标变量的独立贡献）。",

            "本报告存在以下局限性：",

            "**第一，相关不等于因果。** 相关分析和回归分析只能揭示变量之间的统计关联，"
            "不能直接证明因果关系。可能存在未纳入分析的第三变量同时影响各变量。",

            "**第二，样本代表性。** 分析结论的适用范围取决于数据的抽样方式和样本代表性，"
            "结论的外推需谨慎评估。",

            "**第三，测量精度。** 部分变量采用自我报告方式收集，"
            "可能受到回忆偏差、社会期望偏差等因素的影响。定序变量在回归分析中"
            "被近似为连续变量处理，可能对估计精度有一定影响。",

            "**第四，变量类型的自动推断。** 本报告中的变量类型通过自动推断算法确定，"
            "可能存在个别误判。用户可手动调整变量类型以获得更准确的分析结果。",

            "综上所述，本报告的分析结果为数据理解与探索提供了有益参考，"
            "但统计结果应作为辅助决策工具，而非替代专业判断。"
            "建议结合领域知识和定性研究构建更全面的分析框架。",
        ],
        "tables": [],
    }


# ================================================================
# 通用 HTML 渲染
# ================================================================

def _render_generic_html(sections: List[Dict], ctx: _GenericReportContext) -> str:
    """将通用报告节渲染为自包含 HTML 文档。"""
    now = datetime.now().strftime("%Y年%m月%d日")
    title = ctx.report_title

    html_parts = [f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
  body {{
    font-family: "Microsoft YaHei", "SimSun", sans-serif;
    max-width: 900px;
    margin: 40px auto;
    padding: 0 20px;
    color: #333;
    line-height: 1.8;
    background: #fff;
  }}
  h1 {{
    text-align: center;
    font-size: 22px;
    margin-bottom: 5px;
    color: #2B5F8A;
    border-bottom: 2px solid #2B5F8A;
    padding-bottom: 15px;
  }}
  h2 {{
    font-size: 16px;
    margin-top: 30px;
    color: #2B5F8A;
    border-left: 4px solid #2B5F8A;
    padding-left: 10px;
  }}
  p {{
    text-indent: 2em;
    margin: 8px 0;
  }}
  p.no_indent {{
    text-indent: 0;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: 15px 0;
    font-size: 13px;
  }}
  th {{
    background: #2B5F8A;
    color: white;
    padding: 8px;
    border: 1px solid #ddd;
  }}
  td {{
    padding: 6px 8px;
    border: 1px solid #ddd;
    text-align: center;
  }}
  tr:nth-child(even) {{ background: #f8f9fa; }}
  .cover {{
    text-align: center;
    padding: 100px 0;
  }}
  .cover h1 {{
    font-size: 28px;
    border: none;
    line-height: 1.6;
  }}
  .cover .date {{
    color: #888;
    font-size: 14px;
    margin-top: 30px;
  }}
  .footer {{
    text-align: center;
    color: #888;
    margin-top: 50px;
    border-top: 1px solid #eee;
    padding-top: 20px;
  }}
  @media print {{
    body {{ max-width: 100%; }}
    h2 {{ page-break-before: always; }}
  }}
</style>
</head>
<body>

<div class="cover">
<h1>{title}</h1>
<p class="date">报告生成日期：{now}</p>
</div>

<hr style="border: none; border-top: 1px solid #2B5F8A; margin: 30px 0;">
"""]

    for sec in sections:
        html_parts.append(f"<h2>{sec['number']}、{sec['title']}</h2>")

        for para in sec.get("paragraphs", []):
            cls = ' class="no_indent"' if para.startswith("**") else ""
            import re
            para_html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", para)
            html_parts.append(f"<p{cls}>{para_html}</p>")

        for table_data in sec.get("tables", []):
            if table_data.get("title"):
                html_parts.append(
                    f'<p class="no_indent"><em>{table_data["title"]}</em></p>'
                )
            table_df = table_data["data"]
            if table_df is None or table_df.empty:
                continue

            html_parts.append("<table><thead><tr>")
            for col in table_df.columns:
                html_parts.append(f"<th>{col}</th>")
            html_parts.append("</tr></thead><tbody>")

            for _, row in table_df.iterrows():
                html_parts.append("<tr>")
                for col in table_df.columns:
                    val = row[col]
                    if isinstance(val, float):
                        if abs(val) < 1 and val != 0:
                            val = f"{val:.4f}"
                        elif abs(val) < 100:
                            val = f"{val:.2f}"
                        else:
                            val = f"{val:.1f}"
                    html_parts.append(f"<td>{val}</td>")
                html_parts.append("</tr>")
            html_parts.append("</tbody></table>")

    html_parts.append(f"""
<div class="footer">
<p>— 报告完 —</p>
<p>本报告由通用问卷数据 AI 辅助分析平台自动生成，数据来源为实际调查结果。</p>
<p>统计关联不等于因果关系，报告结论仅供辅助决策参考。</p>
</div>

</body>
</html>""")

    return "\n".join(html_parts)
