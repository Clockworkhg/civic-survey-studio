"""统一报告渲染管线。

将 Markdown 报告文本转换为 HTML / DOCX 格式，并处理 AI 返回内容清洗。

管线:
  LLM response content
    → extract_report_content()
    → sanitize_markdown_text()
    → render_markdown_to_html()  /  render_markdown_to_docx()

v0.1.0 Phase 4: 解决 DOCX Markdown 残留、HTML 转义符、JSON 包裹、列表/表格缺失。
"""

from __future__ import annotations

import io
import json
import re
from typing import Any, Dict, List, Optional, Tuple

# ================================================================
# 1. AI 返回内容提取
# ================================================================


def extract_report_content(llm_output: str) -> str:
    """从 LLM 原始输出中提取报告正文。

    处理以下情况:
      1. 纯 Markdown 文本 → 直接返回
      2. JSON 包裹: {"report": "..."} / {"markdown": "..."} / {"content": "..."}
      3. 嵌套 JSON 字符串 → 先反序列化再提取
      4. 被转义的 JSON → 尝试反转义后提取

    Args:
        llm_output: LLM 返回的原始文本

    Returns:
        提取后的 Markdown 报告文本
    """
    if not llm_output or not llm_output.strip():
        return ""

    text = llm_output.strip()

    # Case 1: 纯 Markdown（以 # 开头或以普通文本开头）
    if text.startswith("#") or text.startswith("**") or text.startswith("*"):
        return text
    # 不以 { 或 [ 开头 → 大概率是纯 Markdown
    if not (text.startswith("{") or text.startswith("[")):
        return text

    # Case 2: 尝试 JSON 解析
    try:
        parsed = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        return text  # 不是有效 JSON，直接返回

    # 如果是列表，尝试连接
    if isinstance(parsed, list):
        # 尝试找到第一个含 report/markdown/content 的项
        for item in parsed:
            if isinstance(item, dict):
                content = _extract_from_dict(item)
                if content:
                    return content
        # 列表中没有找到 → 返回原文
        return text

    if isinstance(parsed, dict):
        content = _extract_from_dict(parsed)
        if content:
            return content

    # 如果解析后的对象不包含目标字段，尝试将其序列化后的字符串再次解析
    # （处理双重 JSON 编码的情况）
    if isinstance(parsed, str) and parsed.strip():
        try:
            inner = json.loads(parsed)
            if isinstance(inner, dict):
                content = _extract_from_dict(inner)
                if content:
                    return content
        except (json.JSONDecodeError, ValueError):
            pass

    return text


def _extract_from_dict(obj: Dict[str, Any]) -> str:
    """从 JSON 对象中提取报告内容。"""
    for key in ("report", "markdown", "content", "text", "body", "markdown_report", "analysis_report"):
        val = obj.get(key, "")
        if isinstance(val, str) and val.strip():
            return val.strip()
    # 检查 choices[0].message.content (OpenAI 格式)
    choices = obj.get("choices", [])
    if isinstance(choices, list) and choices:
        msg = choices[0].get("message", {})
        if isinstance(msg, dict):
            content = msg.get("content", "")
            if isinstance(content, str) and content.strip():
                return content.strip()
    return ""


# ================================================================
# 2. Markdown 文本清洗
# ================================================================


def sanitize_markdown_text(text: str) -> str:
    """清洗 Markdown 文本中的常见问题。

    处理:
      - 字面量 \\n → 真正换行
      - \\/ → /
      - 多余的 JSON 转义反斜杠
      - 中英文混合空格规范化
      - 去除首尾空白

    Args:
        text: 待清洗的 Markdown 文本

    Returns:
        清洗后的 Markdown 文本
    """
    if not text:
        return ""

    # 1. 字面量 \n → 真正换行
    text = text.replace("\\n", "\n")

    # 2. \/ → /（JSON 转义残留）
    text = text.replace("\\/", "/")

    # 3. \\" → "（双反斜杠引号）
    text = text.replace('\\"', '"')

    # 4. \\' → '
    text = text.replace("\\'", "'")

    # 5. \\t → \t
    text = text.replace("\\t", "\t")

    # 6. 去除行首尾多余空白（但保留段落间的空行）
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        cleaned.append(line.rstrip())
    text = "\n".join(cleaned)

    # 7. 压缩连续空行（最多保留一个）
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 8. 去除首尾空白
    text = text.strip()

    return text


def clean_html_escapes(html: str) -> str:
    """清理 HTML 中的残留问题。

    处理:
      - \\/ → /
      - 字面量 \\n 在非代码块区域
      - 多余空格

    Args:
        html: HTML 源码字符串

    Returns:
        干净的 HTML 字符串
    """
    if not html:
        return ""

    html = html.replace("\\/", "/")
    html = html.replace("\\n", "\n")
    html = html.replace('\\"', '"')

    return html


def is_probably_json_string(text: str) -> bool:
    """判断文本是否为 JSON 字符串（非 Markdown 正文）。"""
    if not text or not text.strip():
        return False
    t = text.strip()
    return (t.startswith("{") and t.endswith("}")) or (t.startswith("[") and t.endswith("]"))


# ================================================================
# 3. Markdown → HTML 渲染
# ================================================================


def render_markdown_to_html(
    markdown_text: str,
    html_theme: str = "简洁课程作业风",
    report_title: Optional[str] = None,
) -> str:
    """将 Markdown 报告渲染为完整 HTML 页面。

    委托给 html_report_templates.render_html_report，
    但在此之前先执行清洗和内容提取。

    Args:
        markdown_text: Markdown 格式的报告正文
        html_theme: HTML 展示主题
        report_title: 报告标题（用于 <title> 和 <h1>）

    Returns:
        完整的 HTML 文档字符串
    """
    # ── 清洗 ──
    markdown_text = sanitize_markdown_text(markdown_text)
    markdown_text = extract_report_content(markdown_text)

    # ── 委托给现有 HTML 渲染 ──
    from src.html_report_templates import render_html_report as _render_html
    html = _render_html(
        markdown_content=markdown_text,
        html_theme=html_theme,
        report_title=report_title,
    )

    # ── 最终清理 ──
    html = clean_html_escapes(html)

    return html


# ================================================================
# 4. Markdown → DOCX 渲染
# ================================================================


def render_markdown_to_docx(
    markdown_text: str,
    config: Dict[str, Any],
    output_path: Optional[str] = None,
) -> bytes:
    """将 Markdown 报告转为结构化的 DOCX 文件。

    支持:
      - # / ## / ### → Word Heading 1/2/3
      - **粗体** → Word bold run
      - *斜体* → Word italic run
      - - 无序列表 → Word bullet list
      - 1. 有序列表 → Word numbered list
      - Markdown 表格 → Word table
      - 普通段落 → Word normal paragraph
      - 代码块（```...```）→ Word 固定宽度段落
      - 引用（> ...）→ Word 缩进斜体段落

    Args:
        markdown_text: Markdown 格式的报告正文
        config: 分析配置（含 report_title）
        output_path: 可选的文件保存路径（不保存时设为 None）

    Returns:
        .docx 文件字节流
    """
    # ── 清洗 ──
    markdown_text = sanitize_markdown_text(markdown_text)
    markdown_text = extract_report_content(markdown_text)

    # ── 预处理：转义中文星号 ──
    from src.utils import escape_chinese_asterisks
    markdown_text = escape_chinese_asterisks(markdown_text)

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        import logging
        logging.getLogger(__name__).warning("python-docx 未安装，无法生成 DOCX。")
        return b""

    title = config.get("report_title", "数据分析报告")
    doc = Document()

    # ── 默认字体 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ── 标题页 ──
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 逐行解析渲染 ──
    lines = markdown_text.split("\n")
    _render_docx_body(doc, lines)

    # ── 页脚 ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("本报告由 AI 辅助生成 · 统计结果由程序计算 · 文字由大语言模型撰写")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    # ── 保存 ──
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    if output_path:
        with open(output_path, "wb") as f:
            f.write(docx_bytes)

    return docx_bytes


def _render_docx_body(doc, lines: List[str]) -> None:
    """将 Markdown 行渲染到 python-docx Document 对象。

    状态机:
      - normal: 正常段落
      - table: 表格中
      - code_block: 代码块中
      - ul_list: 无序列表中
      - ol_list: 有序列表中
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    i = 0
    n = len(lines)
    current_table = None
    in_code_block = False
    code_lines: List[str] = []

    while i < n:
        line = lines[i]

        # ── 代码块 ──
        if line.startswith("```"):
            if in_code_block:
                # 结束代码块
                _flush_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── 表格 ──
        if "|" in line and line.strip().startswith("|"):
            table_rows = _collect_table_rows(lines, i)
            _render_docx_table(doc, table_rows)
            i += len(table_rows) + 1  # +1 for separator line
            # Skip separator line iterations
            continue

        # ── 无序列表 ──
        ul_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if ul_match:
            items = _collect_list_items(lines, i, r"^(\s*)[-*+]\s+")
            _render_docx_bullet_list(doc, items)
            i += len(items)
            continue

        # ── 有序列表 ──
        ol_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if ol_match:
            items = _collect_list_items(lines, i, r"^(\s*)\d+[.)]\s+")
            _render_docx_numbered_list(doc, items)
            i += len(items)
            continue

        # ── 引用 ──
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # ── 标题 ──
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue

        # ── 水平线 ──
        if line.strip() in ("---", "***", "___", "* * *", "- - -"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # ── 空行 ──
        if line.strip() == "":
            i += 1
            continue

        # ── 普通段落 ──
        p = doc.add_paragraph()
        _render_docx_inline(p, line, bold_pattern=r"\*\*(.+?)\*\*",
                           italic_pattern=r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
        i += 1

    # 清理残留代码块
    if in_code_block and code_lines:
        _flush_code_block(doc, code_lines)


# ── DOCX 辅助 ──

def _render_docx_inline(
    paragraph,
    text: str,
    bold_pattern: str = r"\*\*(.+?)\*\*",
    italic_pattern: str = r"\*(.+?)\*",
) -> None:
    """将带内联格式的文本渲染到段落，处理 **bold** 和 *italic*。

    使用正则分段 → 交替渲染纯文本/加粗/斜体。
    """
    from docx.shared import Pt

    # 组合 pattern: 先匹配 **bold** 再匹配 *italic*
    combined = re.compile(
        r"(\*\*(.+?)\*\*)"   # group 1 = full **text**, group 2 = inner
        r"|"
        r"(\*(.+?)\*)"       # group 3 = full *text*, group 4 = inner
    )

    last_end = 0
    for m in combined.finditer(text):
        # 添加匹配前的纯文本
        prefix = text[last_end:m.start()]
        if prefix:
            paragraph.add_run(prefix)

        if m.group(1):  # Bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # Italic
            run = paragraph.add_run(m.group(4))
            run.italic = True

        last_end = m.end()

    # 添加剩余文本
    suffix = text[last_end:]
    if suffix:
        paragraph.add_run(suffix)


def _collect_table_rows(lines: List[str], start: int) -> List[List[str]]:
    """收集连续的 Markdown 表格行（含表头）。"""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and "|" in line:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(cells)
            i += 1
        else:
            break
    return rows


def _render_docx_table(doc, rows: List[List[str]]) -> None:
    """将 Markdown 表格行渲染为 Word 表格。

    至少需要 3 行: header + separator + data。
    Header 使用加粗字体和灰色背景。
    """
    if len(rows) < 2:
        # 不够组成表格 → 渲染为段落
        for row in rows:
            p = doc.add_paragraph()
            p.add_run(" | ".join(row))
        return

    # Skip separator line
    header = rows[0]
    data_rows = []
    for r in rows[1:]:
        if all(re.match(r"^:?-{2,}:?$", c) for c in r):
            continue  # separator
        data_rows.append(r)

    if not data_rows:
        return

    ncols = max(len(header), max((len(r) for r in data_rows), default=0))
    if ncols == 0:
        return

    # Normalize column count
    header = _pad_row(header, ncols)
    data_rows = [_pad_row(r, ncols) for r in data_rows]

    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn

    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.style = "Table Grid"

    # Header
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(10)
        # Grey background
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8E8E8")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.add_run(cell_text).font.size = Pt(10)

    doc.add_paragraph()  # Spacing after table


def _pad_row(row: List[str], n: int) -> List[str]:
    while len(row) < n:
        row.append("")
    return row[:n]


def _collect_list_items(lines: List[str], start: int, marker_pattern: str) -> List[str]:
    """收集连续的列表项。

    使用提供的 marker_pattern 匹配行，然后剥离 marker 部分提取内容。
    marker_pattern 应包含一个捕获组来匹配列表内容。
    """
    items = []
    i = start
    pat = re.compile(marker_pattern)
    while i < len(lines):
        m = pat.match(lines[i])
        if m:
            # Try group(2) first (if marker_pattern has 2 groups like r"^(\s*)[-*+]\s+(.+)$")
            # Otherwise use group(1)
            if m.lastindex and m.lastindex >= 2:
                items.append(m.group(2))
            elif m.lastindex and m.lastindex >= 1:
                items.append(m.group(1))
            else:
                # Fallback: strip the marker manually
                items.append(re.sub(r"^\s*[-*+]\s+", "", lines[i]))
            i += 1
        else:
            break
    return items


def _render_docx_bullet_list(doc, items: List[str]) -> None:
    """渲染无序列表到 Word。"""
    from docx.shared import Pt
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        _render_docx_inline(p, item)


def _render_docx_numbered_list(doc, items: List[str]) -> None:
    """渲染有序列表到 Word。"""
    from docx.shared import Pt
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.clear()
        _render_docx_inline(p, item)


def _flush_code_block(doc, code_lines: List[str]) -> None:
    """将代码块渲染为 Word 固定宽度段落。"""
    if not code_lines:
        return
    from docx.shared import Pt, RGBColor
    for cl in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(cl)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(60, 60, 60)
